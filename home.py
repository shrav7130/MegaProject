
import streamlit as st
import requests
import tensorflow as tf
import datetime as dt
import os
from apikey import API_KEY
import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np

def get_data(CITY):
    BASE_URL = "http://api.openweathermap.org/data/2.5/weather?"
    url = BASE_URL + "appid=" + API_KEY + "&q=" + CITY

    response = requests.get(url).json()
    json_data = response

    return json_data  # Return the JSON data directly

def predict():
    page_bg_img=    """

        <style>
        [data-testid="stAppViewContainer"]
           {
                background-image: url('https://i.pinimg.com/564x/df/72/16/df721678f11669fc5a6c8199e0ad3871.jpg');
                background-size: cover;
            }
        </style>
        """
    st.markdown(page_bg_img,
        unsafe_allow_html=True
    )
    st.title('Cloud Burst Prediction 🌧️')
    CITY = st.text_input('Enter city name: ')

    if CITY:  # Check if CITY is not empty
        try:
            result = get_data(CITY)

            if result:
                # Extract relevant features from JSON data
                feature_names = ['coord.lat', 'coord.lon', 'main.temp', 'main.feels_like',
                                'main.pressure', 'main.humidity', 'wind.speed', 'wind.deg']

                extracted_features = [get_nested_value(
                    result, name) for name in feature_names]

                # Load the model outside the function to avoid loading it with every prediction
                if 'model' not in st.session_state:
                    st.session_state.model = tf.keras.models.load_model(
                        'api_model')

                model = st.session_state.model

                # Make the prediction
                pred = model.predict([extracted_features])
                pred[0] /= 3

                # Convert wind speed from m/s to km/h and round off to 2 decimal places
                wind_speed_kmh = round(result["wind"]["speed"] * 3.6, 2)

                 # Determine color based on prediction result
                if pred[0][0] < 30.00:  # Low chances
                    color = '#90EE90'
                elif pred[0][0] < 60.00:  # Moderate chances
                    color = 'orange'
                else:  # High chances
                    color = 'red'

                # Display prediction result with dynamically set color
                st.markdown(f'<h1 style="color:{color}; font-size: 35px;">Chances of a Cloud Burst in {CITY} is {pred[0][0]:.2f} %</h1>', unsafe_allow_html=True)

                st.markdown("<br>", unsafe_allow_html=True)


                # Display additional information using a 3x2 grid layout
                st.subheader('Additional Information:')

                # Define a 3x2 grid layout
                row1_col1, row1_col2 = st.columns(2)
                row1_col3, row2_col1 = st.columns(2)
                row2_col2, row2_col3 = st.columns(2)

                # Row 1, Column 1: Temperature
                with row1_col1:
                    st.subheader('Temperature 🌡:')
                    st.markdown(f'<p style="font-size: 25px;"> {result["main"]["temp"]} K</p>', unsafe_allow_html=True)

                # Row 1, Column 2: Humidity
                with row1_col2:
                    st.subheader('Humidity :')
                    st.markdown(f'<p style="font-size:25px;"> {result["main"]["humidity"]}%',unsafe_allow_html=True)

                # Row 1, Column 3: Wind Speed
                with row1_col3:
                    st.subheader('Wind Speed 💨:')
                    st.markdown(f'<p style="font-size:25px;"> {wind_speed_kmh} km/h',unsafe_allow_html=True)

                # Row 2, Column 1: Pressure
                with row2_col1:
                    st.subheader('Pressure :')
                    st.markdown(f'<p style="font-size:25px;">  {result["main"]["pressure"]}Pa',unsafe_allow_html=True)

                # Row 2, Column 2: Wind Degree
                with row2_col2:
                    st.subheader('Wind Degree :')
                    st.markdown(f'<p style="font-size:25px;"> {result["wind"]["deg"]}°',unsafe_allow_html=True)

                # Row 2, Column 3: Latitude and Longitude
                with row2_col3:
                    st.subheader('Latitude and Longitude:')
                    st.markdown(f'<p style="font-size:25px;">Latitude: {result["coord"]["lat"]}',unsafe_allow_html=True)
                    st.markdown(f'<p style="font-size:25px;">Longitude: {result["coord"]["lon"]}',unsafe_allow_html=True)

                # # Define a centered column for the download button
                # st.markdown("<h2 style='text-align: center;'>Download Report</h2>", unsafe_allow_html=True)
                col1, col2, col3 = st.columns([1, 2, 1])



            # Report info
                data = {
                    'Temperature (K)': result["main"]["temp"],
                    'Humidity (%)': result["main"]["humidity"],
                    'Pressure (Pa)': result["main"]["pressure"],
                    'Wind Speed (km/h)': wind_speed_kmh,
                    'Wind Degree (°)': result["wind"]["deg"],
                    'Latitude': result["coord"]["lat"],
                    'Longitude': result["coord"]["lon"],
                    'Chances of Cloud Burst (%)': pred[0][0],
                }

                additional_info_df = pd.DataFrame([data])
                # Create a Word document
                document = Document()
                document.add_heading(f"Weather Report for {CITY}", level=1)

                # Add information to the Word document
                for column, value in data.items():
                    document.add_paragraph(f"{column}: {value}")

                # Save the Word document
                word_file_path = f"{CITY}_weather_report.docx"
                document.save(word_file_path)

                # Download button for the Word document
                with col1:
                    pass

                with col2:
                    st.download_button(
                        label="Download Report",
                        data=open(word_file_path, 'rb').read(),
                        file_name=word_file_path,
                        key="download_report_button"
                    )

                with col3:
                    pass
                # Remove the temporary Word document file
                os.remove(word_file_path)

        except:
            st.error("Please enter a valid city name!")

def get_nested_value(obj, key):
    keys = key.split('.')
    for k in keys:
        if isinstance(obj, dict) and k in obj:
            obj = obj[k]
        elif isinstance(obj, list) and len(obj) > 0:
            obj = obj[int(k)]
        else:
            return None
    return obj

if __name__ == '__main__':
    predict()
